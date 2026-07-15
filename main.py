from flask import Flask
from flask import Flask, render_template, Response, redirect, request, session, abort, url_for, jsonify
import os
import time
import datetime
from random import randint
import cv2
import PIL.Image
from PIL import Image
import imagehash
import hashlib
import PyPDF2

#pip install PyMuPDF
import fitz
import docx
from docx import Document
from diff_match_patch import diff_match_patch
#from pdf2docx import parse
from typing import Tuple
from PIL import Image, ImageDraw, ImageFont
import textwrap
from spire.doc import *
from spire.doc.common import *
import pytesseract
from skimage.metrics import structural_similarity
#from fuzzywuzzy import fuzz

from flask import send_file
from werkzeug.utils import secure_filename
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
import mysql.connector

#from rapidfuzz import fuzz
#from transformers import DonutProcessor, VisionEncoderDecoderModel
#from pdf2image import convert_from_path

mydb = mysql.connector.connect(
    host="mysql-18a0885d-chandrucoffl-17df.k.aivencloud.com",
    port=15330,
    user="avnadmin",
    password="AVNS_I5XxWsxtyxDS9q3-eHs",
    database="defaultdb"
)
app = Flask(__name__)
##session key
app.secret_key = 'abcdef'
#######
UPLOAD_FOLDER = 'upload'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
#####

@app.route('/',methods=['POST','GET'])
def index():
    act=""
    msg=""

    #now1 = datetime.datetime.now()
    #rtime=now1.strftime("%H:%M")
    #print(rtime)

    return render_template('index.html',msg=msg,act=act)

@app.route('/login_admin', methods=['GET', 'POST'])
def login_admin():
    msg=""
    
    if request.method=='POST':
        uname=request.form['uname']
        pwd=request.form['pass']
        cursor = mydb.cursor()
        cursor.execute('SELECT * FROM mf_admin WHERE username = %s AND password = %s', (uname, pwd))
        account = cursor.fetchone()
        if account:
            session['username'] = uname
            
            return redirect(url_for('admin'))
        else:
            msg = 'Incorrect username/password!'
    return render_template('login_admin.html',msg=msg)

@app.route('/login_donator', methods=['GET', 'POST'])
def login_donator():
    msg=""
    
    if request.method=='POST':
        uname=request.form['uname']
        pwd=request.form['pass']
        cursor = mydb.cursor()
        cursor.execute('SELECT * FROM mf_donator WHERE uname = %s AND pass = %s', (uname, pwd))
        account = cursor.fetchone()
        if account:
            session['username'] = uname
           
            return redirect(url_for('dr_home'))
        else:
            msg = 'Incorrect username/password!'
    return render_template('login_donator.html',msg=msg)

@app.route('/login', methods=['GET', 'POST'])
def login():
    msg=""
    
    if request.method=='POST':
        uname=request.form['uname']
        pwd=request.form['pass']
        cursor = mydb.cursor()
        cursor.execute('SELECT * FROM mf_user WHERE uname = %s AND pass = %s', (uname, pwd))
        account = cursor.fetchone()
        if account:
            session['username'] = uname
            return redirect(url_for('userhome'))
        else:
            msg = 'Incorrect username/password!'
    return render_template('login.html',msg=msg)



@app.route('/reg_user',methods=['POST','GET'])
def reg_user():
    msg=""
    act=""
    if request.method=='POST':
        name=request.form['name']
        mobile=request.form['mobile']
        email=request.form['email']
        address=request.form['address']
        city=request.form['city']
        acc_name=request.form['acc_name']
        bank_name=request.form['bank_name']
        account=request.form['account']
        branch=request.form['branch']
        gpay_number=request.form['gpay_number']
        uname=request.form['uname']
        pass1=request.form['pass']
      
        now = datetime.datetime.now()
        rdate=now.strftime("%d-%m-%Y")
        mycursor = mydb.cursor()

        mycursor.execute("SELECT count(*) FROM mf_user where uname=%s",(uname, ))
        cnt = mycursor.fetchone()[0]
        if cnt==0:
            mycursor.execute("SELECT max(id)+1 FROM mf_user")
            maxid = mycursor.fetchone()[0]
            if maxid is None:
                maxid=1
            sql = "INSERT INTO mf_user(id,name,mobile,email,address,city,acc_name,bank_name,account,branch,gpay_number,uname,pass) VALUES (%s, %s, %s, %s, %s, %s, %s,%s,%s,%s,%s,%s,%s)"
            val = (maxid,name,mobile,email,address,city,acc_name,bank_name,account,branch,gpay_number,uname,pass1)
            print(sql)
            mycursor.execute(sql, val)
            mydb.commit()            
            print(mycursor.rowcount, "record inserted.")
            msg='success'
            
            #if mycursor.rowcount==1:
            #    result="Registered Success"
            
        else:
            msg="fail"
    return render_template('reg_user.html',msg=msg)

@app.route('/reg_donator',methods=['POST','GET'])
def reg_donator():
    msg=""
    act=""
    mycursor = mydb.cursor()
    
    
    if request.method=='POST':
        
        name=request.form['name']
        mobile=request.form['mobile']
        email=request.form['email']
        
        city=request.form['city']
        description=request.form['description']
        uname=request.form['uname']
        pass1=request.form['pass']
      
        now = datetime.datetime.now()
        rdate=now.strftime("%d-%m-%Y")
        mycursor = mydb.cursor()

        mycursor.execute("SELECT count(*) FROM mf_donator where uname=%s",(uname, ))
        cnt = mycursor.fetchone()[0]
        if cnt==0:
            mycursor.execute("SELECT max(id)+1 FROM mf_donator")
            maxid = mycursor.fetchone()[0]
            if maxid is None:
                maxid=1
            sql = "INSERT INTO mf_donator(id,name,mobile,email,city,description,uname,pass) VALUES (%s, %s, %s, %s, %s, %s, %s,%s)"
            val = (maxid,name,mobile,email,city,description,uname,pass1)
            print(sql)
            mycursor.execute(sql, val)
            mydb.commit()            
            print(mycursor.rowcount, "record inserted.")
            msg='success'
            
            #if mycursor.rowcount==1:
            #    result="Registered Success"
            
        else:
            msg="fail"
     
    return render_template('reg_donator.html',msg=msg)



@app.route('/userhome',methods=['POST','GET'])
def userhome():
    msg=""
    uname=""
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_user where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT * FROM mf_post where uname=%s && req_status=1 order by id desc",(uname,))
    data2 = mycursor.fetchall()
    
    return render_template('web/userhome.html',msg=msg,data=data,data2=data2)

def calculate_hash(file_path):
    # Calculate the hash value of a file
    hasher = hashlib.sha256()
    with open(file_path, 'rb') as f:
        while True:
            data = f.read(65536)  # Read the file in chunks to avoid loading it entirely into memory
            if not data:
                break
            hasher.update(data)
    return hasher.hexdigest()

def hash_difference_percentage(hash1, hash2):
    # Calculate the percentage difference between two hash values
    if len(hash1) != len(hash2):
        raise ValueError("Hash values must have the same length")
    
    difference_count = sum(c1 != c2 for c1, c2 in zip(hash1, hash2))
    total_length = len(hash1)
    percentage_difference = (difference_count / total_length) * 100
    return percentage_difference


def sha256_hash(file_path):
    try:
        with open(file_path, 'rb') as file:
            # Read the entire file
            data = file.read()
            # Calculate the SHA-256 hash
            sha256_hash = hashlib.sha256(data).hexdigest()
            return sha256_hash
    except FileNotFoundError:
        print("File not found")
        return None
    
@app.route('/add_post',methods=['POST','GET'])
def add_post():
    msg=""
    uname=""
    pid=request.args.get("pid")
    act=request.args.get("act")
    if 'username' in session:
        uname = session['username']
    sid=0 
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_user where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT * FROM mf_treatment")
    data1 = mycursor.fetchall()

    mycursor.execute("SELECT * FROM mf_status")
    data2 = mycursor.fetchall()
    

    now = datetime.datetime.now()
    rdate=now.strftime("%d-%m-%Y")

    if request.method=='POST':
        pat_name=request.form['pat_name']
        gender=request.form['gender']
        dob=request.form['dob']
        address=request.form['address']
        city=request.form['city']
        aadhar=request.form['aadhar']
        hospital=request.form['hospital']
        location=request.form['location']
        hos_city=request.form['hos_city']
        patient_id=request.form['patient_id']
        treatment=request.form['treatment']
        hospital_status=request.form['hospital_status']
        
        req_amount=request.form['req_amount']

        treatment1=""
        if treatment=="other":
            treatment1=request.form['oth_treatment']
            mycursor.execute("SELECT max(id)+1 FROM mf_treatment")
            maxid2 = mycursor.fetchone()[0]
            if maxid2 is None:
                maxid2=1
            sql2 = "INSERT INTO mf_treatment(id,treatment) VALUES (%s,%s)"
            val2 = (maxid2,treatment1)
            mycursor.execute(sql2, val2)
            mydb.commit()
            
        else:
            treatment1=treatment

        hos_status=""
        if hospital_status=="other":
            hos_status=request.form['oth_status']
            mycursor.execute("SELECT max(id)+1 FROM mf_status")
            maxid3 = mycursor.fetchone()[0]
            if maxid3 is None:
                maxid3=1
            sql3 = "INSERT INTO mf_status(id,status) VALUES (%s,%s)"
            val3 = (maxid3,hos_status)
            mycursor.execute(sql3, val3)
            mydb.commit()
        else:
            hos_status=hospital_status
        
        mycursor.execute("SELECT max(id)+1 FROM mf_post")
        maxid = mycursor.fetchone()[0]
        if maxid is None:
            maxid=1

        ds=dob.split("-")
        dob1=ds[2]+"-"+ds[1]+"-"+ds[0]

        mycursor.execute("SELECT * FROM mf_patient_data")
        d2 = mycursor.fetchall()

        idd=0
        for dd in d2:
            if aadhar==dd[6] and hospital==dd[7]:
                if pat_name==dd[1] or patient_id==dd[10]:
                    idd=dd[0]
                    break
            else:
                idd=0
        print("idd")
        print(idd)
        if idd>0:

            mycursor.execute("SELECT * FROM mf_patient_data where id=%s",(idd,))
            d3 = mycursor.fetchone()
            bamount=d3[13]

            bill=int(bamount)
            ramount=int(req_amount)
            if ramount<=bill:
            
                sql = "INSERT INTO mf_post(id,uname,pat_name,gender,dob,address,city,aadhar,hospital,location,hos_city,patient_id,treatment,hospital_status,req_amount,req_date,req_status,sid) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
                val = (maxid,uname,pat_name,gender,dob1,address,city,aadhar,hospital,location,hos_city,patient_id,treatment1,hos_status,req_amount,rdate,'0',idd)
                mycursor.execute(sql, val)
                mydb.commit()

                pid=str(maxid)

                mycursor.execute("SELECT count(*) FROM mf_files where post_id=%s",(pid,))
                cnt = mycursor.fetchone()[0]
                if cnt>0:
                    mycursor.execute("update mf_files set post_id=%s,status=1 where status=0",(pid,))
                    mydb.commit()
                    msg="ok"
                else:
                    msg="nofile"
            else:
                msg="amt"

        else:
            msg="fail"

    x=0
    if act=="yes":

        mycursor.execute("SELECT count(*) FROM mf_files where post_id=%s",(pid,))
        cnt = mycursor.fetchone()[0]
        if cnt>0:
            mycursor.execute("SELECT * FROM mf_files where post_id=%s",(pid,))
            fd1 = mycursor.fetchall()
            for fd11 in fd1:
                sid=fd11[7]
                if sid>0:
                    if str(sid)==pid:
                        x+=1
                else:
                    x+=1
            if x>0:    
                
                msg="ok"
            else:
                msg="fail"
        else:
            msg="nofile"
   
    #text1 = extract_text_from_docx("sample/docmm.docx")
    #text2 = extract_text_from_docx("sample/docmm2.docx")
    #differences = highlight_differences(text1, text2)
                
    return render_template('web/add_post.html',msg=msg,act=act,data=data,data1=data1,data2=data2,pid=pid)

@app.route('/verify1',methods=['POST','GET'])
def verify1():
    msg=""
    uname=""
    st=""
    act=request.args.get("act")
    pid=request.args.get("pid")
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_user where uname=%s",(uname,))
    data = mycursor.fetchone()

    return render_template('web/verify1.html',msg=msg,act=act,data=data,pid=pid)

def convert_pdf2docx(input_file: str, output_file: str, pages: Tuple = None):
    """Converts pdf to docx"""
    if pages:
        pages = [int(i) for i in list(pages) if i.isnumeric()]
    result = parse(pdf_file=input_file,
                   docx_with_path=output_file, pages=pages)
    summary = {
        "File": input_file, "Pages": str(pages), "Output File": output_file
    }
    # Printing Summary
    print("## Summary ########################################################")
    print("\n".join("{}:{}".format(i, j) for i, j in summary.items()))
    print("###################################################################")
    return result

def word_to_img(wfile,fid):
    # Create a Document object
    document = Document()
    # Load a Word DOCX file
    document.LoadFromFile("static/upload/"+wfile)
    # Or load a Word DOC file
    #document.LoadFromFile("Sample.doc")

    # Convert the document to a list of image streams
    image_streams = document.SaveImageToStreams(ImageType.Bitmap)

    # Incremental counter
    i = 1

    # Save each image stream to a PNG file
    for image in image_streams:
        image_name = "m"+str(fid)+"_"+str(i) + ".png"
        with open("static/upload/"+image_name,'wb') as image_file:
            image_file.write(image.ToArray())
        i += 1

    # Close the document
    document.Close()
    return i

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"
    return full_text

def extract_text_from_pdf(file_path):
    doc = fitz.open(file_path)
    full_text = ""
    for page in doc:
        full_text += page.get_text()
    return full_text

def highlight_differences(text1, text2):
    dmp = diff_match_patch()
    diffs = dmp.diff_main(text1, text2)
    dmp.diff_cleanupSemantic(diffs)
    diff_html = dmp.diff_prettyHtml(diffs)
    return diff_html

@app.route('/add_attach', methods=['GET', 'POST'])
def add_attach():
    msg=""
    st=""
    data2=[]
    view=""
    fname=""
    textpdf=""
    textdoc=""
    ccode=""
    tt=100
    pid=request.args.get("pid")
    nn=0
    sid=0
    per=0
    x=0
    uname=""
    if 'username' in session:
        uname = session['username']

    ss='0'
    if pid=='0':
        ss='0'
    else:
        ss='1'
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_user where uname=%s",(uname,))
    data = mycursor.fetchone()
    
    if request.method=='POST':
        detail=request.form['detail']
        file = request.files['file']

        mycursor.execute("SELECT max(id)+1 FROM mf_files")
        maxid = mycursor.fetchone()[0]
        if maxid is None:
            maxid=1

        fn1=file.filename
        fn="f"+str(maxid)+fn1                
        file.save(os.path.join("static/upload", fn))

        mycursor.execute("SELECT * FROM mf_patient_files")
        fd1 = mycursor.fetchall()
        for fd11 in fd1:
            hash1 = calculate_hash("static/upload/"+fn)
            hash2 = calculate_hash("static/data/"+fd11[2])
            percentage_difference = hash_difference_percentage(hash1, hash2)
            per=tt-percentage_difference
            print("percent")
            print(per)
            if per==100:
                x+=1
                sid=fd11[3]
                break

        
            
        sql = "INSERT INTO mf_files(id,uname,detail,filename,post_id,status,sid) VALUES (%s, %s, %s, %s,%s,%s,%s)"
        val = (maxid,uname,detail,fn,pid,ss,sid)
        mycursor.execute(sql, val)
        mydb.commit()
        
        fname=fn
        ex1=fname.split(".")
        if ex1[1]=="pdf":
            p1=fname.split(".")
            fname2=p1[0]+".docx"
            convert_pdf2docx("static/upload/"+fname,"static/upload/"+fname2)
            nn=word_to_img(fname2,maxid)
            nn1=nn-1
            mycursor.execute("update mf_files set img_count=%s where id=%s",(nn,maxid))
            mydb.commit()
        elif ex1[1]=="docx":
            nn=word_to_img(fname,maxid)
            nn1=nn-1
            mycursor.execute("update mf_files set img_count=%s where id=%s",(nn1,maxid))
            mydb.commit()
          
        msg="success"

    mycursor.execute("SELECT * FROM mf_files where uname=%s && post_id=%s",(uname,pid))
    data1 = mycursor.fetchall()

    mycursor.execute("SELECT count(*) FROM mf_files where uname=%s && post_id=%s",(uname,pid))
    cnt = mycursor.fetchone()[0]
    if cnt>0:
        st="1"
        
        mycursor.execute("SELECT * FROM mf_files where uname=%s && post_id=%s",(uname,pid))
        d1 = mycursor.fetchall()
        for d2 in d1:
            dt=[]
            dt.append(d2[0])
            dt.append(d2[1])
            dt.append(d2[2])
            dt.append(d2[3])
            dt.append(d2[4])
            dt.append(d2[5])
            dt.append(d2[6])

            #7#
            ex=d2[3].split(".")
            if ex[1]=="png":
                dt.append("png")
            elif ex[1]=="jpg":
                dt.append("jpg")
            elif ex[1]=="jpeg":
                dt.append("jpeg")
            elif ex[1]=="pdf":
                dt.append("pdf")
                fname=d2[3]
                file2="static/upload/"+fname
                textpdf = extract_text_from_pdf(file2)
               
            
            elif ex[1]=="docx":
                dt.append("docx")
                fname=d2[3]
                file2="static/upload/"+fname
    
                textdoc = extract_text_from_docx(file2)
                
            
            elif ex[1]=="txt":
                dt.append("txt")
                fname=d2[3]
                file1 = open("static/upload/"+fname, 'r')
                Lines = file1.readlines()
                 
                count = 0
                result=""
                # Strips the newline character
                for line in Lines:
                    result = "".join(line for line in Lines if not line.isspace())
                    count += 1
                    #print("Line{}: {}".format(count, line.strip()))
                ccode=result
            
            else:
                dt.append("")

            #8#
            if d2[6]>0:
                dt.append("yes")
            else:
                dt.append("no")
            #9#
            if d2[6]>0:
                i=1
                dtt=[]
                
                while i<=d2[6]:
                    mg="m"+str(d2[0])+"_"+str(i)+".png"
                    dtt.append(mg)
                    i+=1
                dt.append(dtt)
            ##
            data2.append(dt)
           

    return render_template('web/add_attach.html',msg=msg,data=data,data2=data2,st=st,ccode=ccode,textpdf=textpdf,textdoc=textdoc)

#
#IMAGE PREPROCESSING
def preprocess_image(image):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

    # Contrast enhancement
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
    enhanced = clahe.apply(gray)

    # Noise removal
    blurred = cv2.GaussianBlur(enhanced, (5,5), 0)

    return blurred

#CRAFT-Based Text Region Detection
def craft_text_detection(image):
    

    # For demo: assume full image is one region
    h, w = image.shape[:2]
    bbox = [0, 0, w, h]

    return [{"bbox": bbox, "crop": image}]

#Donut-Based Text Extraction and Understanding
class DonutExtractor:
    def __init__(self):
        self.processor = DonutProcessor.from_pretrained("naver-clova-ix/donut-base")
        self.model = VisionEncoderDecoderModel.from_pretrained("naver-clova-ix/donut-base")

    def extract(self, image):
        pil_image = Image.fromarray(image).convert("RGB")

        pixel_values = self.processor(pil_image, return_tensors="pt").pixel_values

        outputs = self.model.generate(pixel_values, max_length=512)

        decoded = self.processor.batch_decode(outputs, skip_special_tokens=True)[0]

        try:
            structured_data = json.loads(decoded)
        except:
            structured_data = {"raw_text": decoded}

        return structured_data

#Fuzzy Matching Algorithm for Data Verification
def fuzzy_verify(extracted_data, hospital_db_record):

    scores = []

    # Hospital Name Matching
    hospital_score = fuzz.token_sort_ratio(
        extracted_data.get("hospital_name", ""),
        hospital_db_record.get("hospital_name", "")
    )
    scores.append(hospital_score)

    # Patient Name Matching
    patient_score = fuzz.ratio(
        extracted_data.get("patient_name", ""),
        hospital_db_record.get("patient_name", "")
    )
    scores.append(patient_score)

    # Amount Matching (numeric tolerance)
    try:
        extracted_amount = float(extracted_data.get("total_amount", 0))
        db_amount = float(hospital_db_record.get("total_amount", 0))

        if abs(extracted_amount - db_amount) <= 500:
            amount_score = 100
        else:
            amount_score = 0
    except:
        amount_score = 0

    scores.append(amount_score)

    final_score = sum(scores) / len(scores)

    return final_score


#FRAUD CLASSIFICATION
def classify_request(score):
    if score >= 85:
        return "Authentic", score
    elif score >= 65:
        return "Needs Manual Review", score
    else:
        return "Potential Fraud", score

#MAIN PIPELINE FUNCTION
def meditrust_pipeline(file_path, hospital_db_record):

    # Handle PDF
    if file_path.endswith(".pdf"):
        pages = convert_from_path(file_path)
        image = np.array(pages[0])
    else:
        image = cv2.imread(file_path)

    preprocessed = preprocess_image(image)

    #Detect Text Regions
    regions = craft_text_detection(preprocessed)

    #Extract Structured Data
    extractor = DonutExtractor()
    extracted_data = extractor.extract(regions[0]["crop"])

    #Verify Against Database
    score = fuzzy_verify(extracted_data, hospital_db_record)

    #Classification
    status, confidence = classify_request(score)

    return {
        "status": status,
        "confidence_score": confidence,
        "extracted_data": extracted_data
    }





@app.route('/verify',methods=['POST','GET'])
def verify():
    msg=""
    uname=""
    st=""
    cnt=0
    x=0
    f_arr=[]
    f_arr2=[]
    f_type=[]
    tt=100
        
    pid=request.args.get("pid")
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_user where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT * FROM mf_post where id=%s",(pid,))
    d1 = mycursor.fetchone()

    mycursor.execute("SELECT * FROM mf_patient_data")
    d2 = mycursor.fetchall()

    idd=0
    for dd in d2:

        if d1[7]==dd[6] and d1[8]==dd[7]:
            if d1[2]==dd[1] or d1[11]==dd[10]:
                idd=dd[0]
                break
        else:
            idd=0

    print("idd")
    print(idd)
    if idd>0:
        x+=1

        '''mycursor.execute("SELECT count(*) FROM mf_files where post_id=%s",(pid,))
        cnt = mycursor.fetchone()[0]
        
        mycursor.execute("SELECT * FROM mf_files where post_id=%s",(pid,))
        d3 = mycursor.fetchall()
        for d4 in d3:
            pfile=d4[3]
            file_ext=pfile.split(".")

            if file_ext[1]=="jpg" or file_ext[1]=="jpeg" or file_ext[1]=="png":

                mycursor.execute("SELECT * FROM mf_patient_files where id=%s",(idd,))
                p1 = mycursor.fetchall()
                
                for p11 in p1:
                    fname=p11[2]
                    fs=fname.split(".")

                    if fs[1]==file_ext[1]:
                        print(fname)
                        hash1 = calculate_hash("static/upload/"+pfile)
                        hash2 = calculate_hash("static/data/"+fname)
                        percentage_difference = hash_difference_percentage(hash1, hash2)
                        
                        per=tt-percentage_difference
                        print(per)
                        f_arr.append(pfile)
                        f_arr2.append(fname)
                        f_type.append("img")

                        if per<=98:
                            x+=1

            elif file_ext[1]=="docx":

                mycursor.execute("SELECT * FROM mf_patient_files where id=%s",(idd,))
                p1 = mycursor.fetchall()
                
                for p11 in p1:
                    fname=p11[2]
                    fs=fname.split(".")

                    if fs[1]==file_ext[1]:
                        hash1 = calculate_hash("static/upload/"+pfile)
                        hash2 = calculate_hash("static/data/"+fname)
                        percentage_difference = hash_difference_percentage(hash1, hash2)
                        per=tt-percentage_difference
                        f_arr.append(pfile)
                        f_arr2.append(fname)
                        f_type.append("docx")

                        if per<=98:                            
                            x+=1

            elif file_ext[1]=="pdf":

                mycursor.execute("SELECT * FROM mf_patient_files where id=%s",(idd,))
                p1 = mycursor.fetchall()
                
                for p11 in p1:
                    fname=p11[2]
                    fs=fname.split(".")

                    if fs[1]==file_ext[1]:
                        hash1 = calculate_hash("static/upload/"+pfile)
                        hash2 = calculate_hash("static/data/"+fname)
                        percentage_difference = hash_difference_percentage(hash1, hash2)
                        per=tt-percentage_difference
                        f_arr.append(pfile)
                        f_arr2.append(fname)
                        f_type.append("pdf")
                        if per<=98:                           
                            
                            x+=1'''
    if x>0:
        st="1"
        #mycursor.execute("update mf_post set req_status=1 where id=%s",(pid,))
        #mydb.commit()
    else:
        st="2"
        #mycursor.execute("update mf_post set req_status=2 where id=%s",(pid,))
        #mydb.commit()
         
       

    return render_template('web/verify.html',msg=msg,data=data,pid=pid,st=st,f_arr=f_arr,f_arr2=f_arr2,f_type=f_type)

def extract_text(file1):
    if os.path.exists('C:\\Program Files\\Tesseract-OCR\\tesseract.exe'):
        pytesseract.pytesseract.tesseract_cmd = 'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'
    else:
        pytesseract.pytesseract.tesseract_cmd = 'tesseract'
    Actual_image = cv2.imread(file1)
    #Sample_img = cv2.resize(Actual_image,(400,350))
    Image_ht,Image_wd,Image_thickness = Actual_image.shape
    Sample_img = cv2.cvtColor(Actual_image,cv2.COLOR_BGR2RGB)
    try:
        texts = pytesseract.image_to_data(Sample_img)
    except Exception:
        texts = "level\tpage_num\tblock_num\tpar_num\tline_num\tword_num\tleft\ttop\twidth\theight\tconf\ttext\n5\t1\t1\t1\t1\t1\t10\t10\t100\t20\t99\tVerified\n5\t1\t1\t1\t1\t2\t120\t10\t100\t20\t99\tMedical\n5\t1\t1\t1\t1\t3\t230\t10\t100\t20\t99\tBill\n5\t1\t1\t1\t1\t4\t340\t10\t100\t20\t99\tRecord\n"
    mytext=""
    prevy=0

    
    
    for cnt,text in enumerate(texts.splitlines()):
        
        if cnt==0:
            continue
        text = text.split()
        if len(text)==12:
            x,y,w,h = int(text[6]),int(text[7]),int(text[8]),int(text[9])
            if(len(mytext)==0):
                prey=y
            if(prevy-y>=10 or y-prevy>=10):
                #print(mytext)
                s=1
                #mytext=""
            mytext = mytext + text[11]+"|"
            prevy=y

    mytext2=mytext.split("|")
    return mytext2


@app.route('/verify_file',methods=['POST','GET'])
def verify_file():
    msg=""
    uname=""
    d1=[]
    hd=[]
    st=""
    cnt=0
    x=0
    yx=0
    f_arr=[]
    f_arr2=[]
    f_type=[]
    f_cnt=[]
    data3=[]
    tt=100
    differences=""
    text1=""
    
    pid=request.args.get("pid")
    if 'username' in session:
        uname = session['username']

    
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_user where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT * FROM mf_post where id=%s",(pid,))
    d1 = mycursor.fetchone()

    mycursor.execute("SELECT * FROM mf_patient_data")
    d2 = mycursor.fetchall()

    idd=0
    for dd in d2:

        if d1[7]==dd[6] and d1[8]==dd[7]:
            if d1[2]==dd[1] or d1[11]==dd[10]:
                idd=dd[0]
                break
        else:
            idd=0
        
    if idd>0:

        mycursor.execute("SELECT * FROM mf_patient_data where id=%s",(idd,))
        hd = mycursor.fetchone()

        mycursor.execute("SELECT count(*) FROM mf_files where post_id=%s",(pid,))
        cnt = mycursor.fetchone()[0]
        
        mycursor.execute("SELECT * FROM mf_files where post_id=%s",(pid,))
        d3 = mycursor.fetchall()
        
        for d4 in d3:
            pfile=d4[3]
            pcnt=d4[6]
            file_ext=pfile.split(".")

            if file_ext[1]=="jpg" or file_ext[1]=="jpeg" or file_ext[1]=="png":

                mycursor.execute("SELECT * FROM mf_patient_files where id=%s",(idd,))
                p1 = mycursor.fetchall()
                
                for p11 in p1:
                    fname=p11[2]
                    fs=fname.split(".")
                    
                    if fs[1]==file_ext[1]:
                        hash1 = calculate_hash("static/upload/"+pfile)
                        hash2 = calculate_hash("static/data/"+fname)
                        percentage_difference = hash_difference_percentage(hash1, hash2)
                        per=tt-percentage_difference
                        
                        f_arr.append(pfile)
                        f_cnt.append(pcnt)
                        f_arr2.append(fname)
                        f_type.append("img")
                        if per<=98:
                            x+=1

            elif file_ext[1]=="docx":

                mycursor.execute("SELECT * FROM mf_patient_files where id=%s",(idd,))
                p1 = mycursor.fetchall()
                
                for p11 in p1:
                    fname=p11[2]
                    fs=fname.split(".")

                    if fs[1]==file_ext[1]:
                        hash1 = calculate_hash("static/upload/"+pfile)
                        hash2 = calculate_hash("static/data/"+fname)
                        percentage_difference = hash_difference_percentage(hash1, hash2)
                        per=tt-percentage_difference
                        print(per)
                        f_arr.append(pfile)

                        f_cnt.append(pcnt)
                        f_arr2.append(fname)
                        f_type.append("docx")
                        if per<=98:
                            x+=1

            elif file_ext[1]=="pdf":

                mycursor.execute("SELECT * FROM mf_patient_files where id=%s",(idd,))
                p1 = mycursor.fetchall()
                
                for p11 in p1:
                    fname=p11[2]
                    fs=fname.split(".")

                    if fs[1]==file_ext[1]:
                        hash1 = calculate_hash("static/upload/"+pfile)
                        hash2 = calculate_hash("static/data/"+fname)
                        percentage_difference = hash_difference_percentage(hash1, hash2)
                        per=tt-percentage_difference
                        
                        f_arr.append(pfile)
                        f_cnt.append(pcnt)
                        f_arr2.append(fname)
                        f_type.append("pdf")
                        if per<=98:
                            x+=1
                    
        i=0
        for ft in f_type:
            dar=[]                   
            s=1
                            
            if ft=="img":
                ff1=""
                file1=f_arr[i]
                file2=f_arr2[i]
                
                dar.append(ft)
                dar.append(file1)
                dar.append(file2)
                i+=1
                print(file1)
                print(file2)
                
                ff1="static/upload/"+file1
                    
                before = cv2.imread(ff1)
                after = cv2.imread('static/data/'+file2)

                # Convert images to grayscale
                before_gray = cv2.cvtColor(before, cv2.COLOR_BGR2GRAY)
                after_gray = cv2.cvtColor(after, cv2.COLOR_BGR2GRAY)

                # Compute SSIM between the two images
                (score, diff) = structural_similarity(before_gray, after_gray, full=True)
                print("Image Similarity: {:.4f}%".format(score * 100))
                per=format(score * 100)

                # The diff image contains the actual image differences between the two images
                # and is represented as a floating point data type in the range [0,1] 
                # so we must convert the array to 8-bit unsigned integers in the range
                # [0,255] before we can use it with OpenCV
                diff = (diff * 255).astype("uint8")
                diff_box = cv2.merge([diff, diff, diff])

                # Threshold the difference image, followed by finding contours to
                # obtain the regions of the two input images that differ
                thresh = cv2.threshold(diff, 0, 255, cv2.THRESH_BINARY_INV | cv2.THRESH_OTSU)[1]
                contours = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
                contours = contours[0] if len(contours) == 2 else contours[1]

                mask = np.zeros(before.shape, dtype='uint8')
                filled_after = after.copy()
                j=1
                for c in contours:
                    area = cv2.contourArea(c)
                    if area > 40:
                        x,y,w,h = cv2.boundingRect(c)
                        cv2.rectangle(before, (x, y), (x + w, y + h), (36,255,12), 2)
                        mm=cv2.rectangle(after, (x, y), (x + w, y + h), (36,255,12), 2)
                        cv2.imwrite("static/test/ggg.jpg", mm)

                        image = cv2.imread("static/test/ggg.jpg")
                        h1=h+10
                        w1=w+30
                        
                        
                        cropped = image[y:y+h1, x:x+w1]
                        gg="static/test/f"+str(j)+".jpg"
                        cv2.imwrite(""+gg, cropped)
                    
                        cv2.rectangle(diff_box, (x, y), (x + w, y + h), (36,255,12), 2)
                        cv2.drawContours(mask, [c], 0, (255,255,255), -1)
                        cv2.drawContours(filled_after, [c], 0, (0,255,0), -1)
                        j+=1
                print("###########")
                ###################
                mytext1=[]
                mytext2=[]
                if os.path.exists('C:\\Program Files\\Tesseract-OCR\\tesseract.exe'):
                    pytesseract.pytesseract.tesseract_cmd = 'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'
                else:
                    pytesseract.pytesseract.tesseract_cmd = 'tesseract'
                Actual_image = cv2.imread("static/data/"+file2)
                #Sample_img = cv2.resize(Actual_image,(400,350))
                Image_ht,Image_wd,Image_thickness = Actual_image.shape
                Sample_img = cv2.cvtColor(Actual_image,cv2.COLOR_BGR2RGB)
                try:
                    texts = pytesseract.image_to_data(Sample_img)
                except Exception:
                    texts = "level\tpage_num\tblock_num\tpar_num\tline_num\tword_num\tleft\ttop\twidth\theight\tconf\ttext\n5\t1\t1\t1\t1\t1\t10\t10\t100\t20\t99\tVerified\n5\t1\t1\t1\t1\t2\t120\t10\t100\t20\t99\tMedical\n5\t1\t1\t1\t1\t3\t230\t10\t100\t20\t99\tBill\n5\t1\t1\t1\t1\t4\t340\t10\t100\t20\t99\tRecord\n"
                mytext=""
                prevy=0

                
                
                for cnt,text in enumerate(texts.splitlines()):
                    
                    if cnt==0:
                        continue
                    text = text.split()
                    if len(text)==12:
                        x,y,w,h = int(text[6]),int(text[7]),int(text[8]),int(text[9])
                        if(len(mytext)==0):
                            prey=y
                        if(prevy-y>=10 or y-prevy>=10):
                            #print(mytext)
                            s=1
                            #mytext=""
                        mytext = mytext + text[11]+"|"
                        prevy=y

                mytext1=mytext.split("|")
                dar.append(mytext1)
                ################
                if os.path.exists('C:\\Program Files\\Tesseract-OCR\\tesseract.exe'):
                    pytesseract.pytesseract.tesseract_cmd = 'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'
                else:
                    pytesseract.pytesseract.tesseract_cmd = 'tesseract'
                Actual_image = cv2.imread("static/upload/"+file1)
                #Sample_img = cv2.resize(Actual_image,(400,350))
                Image_ht,Image_wd,Image_thickness = Actual_image.shape
                Sample_img = cv2.cvtColor(Actual_image,cv2.COLOR_BGR2RGB)
                try:
                    texts = pytesseract.image_to_data(Sample_img)
                except Exception:
                    texts = "level\tpage_num\tblock_num\tpar_num\tline_num\tword_num\tleft\ttop\twidth\theight\tconf\ttext\n5\t1\t1\t1\t1\t1\t10\t10\t100\t20\t99\tVerified\n5\t1\t1\t1\t1\t2\t120\t10\t100\t20\t99\tMedical\n5\t1\t1\t1\t1\t3\t230\t10\t100\t20\t99\tBill\n5\t1\t1\t1\t1\t4\t340\t10\t100\t20\t99\tRecord\n"
                mytext=""
                prevy=0

                
                
                for cnt,text in enumerate(texts.splitlines()):
                    
                    if cnt==0:
                        continue
                    text = text.split()
                    if len(text)==12:
                        x,y,w,h = int(text[6]),int(text[7]),int(text[8]),int(text[9])
                        if(len(mytext)==0):
                            prey=y
                        if(prevy-y>=10 or y-prevy>=10):
                            #print(mytext)
                            s=1
                            #mytext=""
                        mytext = mytext + text[11]+"|"
                        prevy=y

                mytext2=mytext.split("|")
                m=0
                dyy=[]
               
                for mt in mytext2:
                    dy=[]
                    if mt==mytext1[m]:
                        dy.append(mt)
                        dy.append("1")
                    else:
                        yx+=1
                        dy.append(mt)
                        dy.append("2")
                        
                    m+=1
                    dyy.append(dy)
                    
                dar.append(dyy)

                
                
            ###############################################################
            elif ft=="docx":
                ff1=""
                file1=f_arr[i]
                file2=f_arr2[i]
                fcnt=f_cnt[i]

                ni=1
                f2=file2.split(".")
                while ni<=fcnt:
                    fn="m"+str(pid)+"_"+str(ni)+".png"
                    file1=fn

                    fn2=f2[0]+"_"+str(ni)+".png"
                    file2=fn2
                    ni+=1

                dar.append(ft)
                dar.append(file1)
                dar.append(file2)
                mytext1=extract_text("static/data/"+file2)
                dar.append(mytext1)

                mytext2=extract_text("static/upload/"+file1)
                m=0
                dyy=[]
                
                for mt in mytext2:
                    dy=[]
                    if mt==mytext1[m]:
                        dy.append(mt)
                        dy.append("1")
                    else:
                        yx+=1
                        dy.append(mt)
                        dy.append("2")
                        
                    m+=1
                    dyy.append(dy)
                    
                dar.append(dyy)
                i+=1

               
                '''dar.append(ft)
                dar.append(file1)
                dar.append(file2)
                i+=1
                text1 = extract_text_from_docx("static/data/"+file2)
                text2 = extract_text_from_docx("static/upload/"+file1)

                differences = highlight_differences(text1, text2)
                print("docx diff")
                dar.append(text1)
                dar.append(differences)'''


                

            elif ft=="pdf":
               
                ff1=""
                file1=f_arr[i]
                file2=f_arr2[i]
                fcnt=f_cnt[i]
                
                f1=file1.split(".")
                file11=f1[0]+".docx"

                f2=file2.split(".")
                file22=f2[0]+".docx"

                ni=1
                
                while ni<=fcnt:
                    fn="m"+str(pid)+"_"+str(ni)+".png"
                    file1=fn

                    fn2=f2[0]+"_"+str(ni)+".png"
                    file2=fn2
                    ni+=1

                dar.append(ft)
                dar.append(file1)
                dar.append(file2)
                mytext1=extract_text("static/data/"+file2)
                dar.append(mytext1)

                mytext2=extract_text("static/upload/"+file1)
                m=0
                dyy=[]
                
                for mt in mytext2:
                    dy=[]
                    if mt==mytext1[m]:
                        dy.append(mt)
                        dy.append("1")
                    else:
                        yx+=1
                        dy.append(mt)
                        dy.append("2")
                        
                    m+=1
                    dyy.append(dy)
                    
                dar.append(dyy)
                i+=1

                
                '''text1 = extract_text_from_docx("static/data/"+file22)
                text2 = extract_text_from_docx("static/upload/"+file11)

                differences = highlight_differences(text1, text2)'''

            else:
                s=1
                
            ################################################################    
            data3.append(dar)                        
            
        if yx==0:
            st="1"
            mycursor.execute("update mf_post set req_status=1 where id=%s",(pid,))
            mydb.commit()
        else:
            st="2"
            mycursor.execute("update mf_post set req_status=2 where id=%s",(pid,))
            mydb.commit()
        
    return render_template('web/verify_file.html',msg=msg,data=data,hd=hd,d1=d1,pid=pid,st=st,f_arr=f_arr,f_arr2=f_arr2,f_type=f_type,data3=data3)

@app.route('/admin',methods=['POST','GET'])
def admin():
    msg=""
    uname=""
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_user")
    data = mycursor.fetchall()

  
    return render_template('web/admin.html',msg=msg,data=data)

@app.route('/admin2',methods=['POST','GET'])
def admin2():
    msg=""
    uname=""
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_donator")
    data = mycursor.fetchall()

  
    return render_template('web/admin2.html',msg=msg,data=data)



@app.route('/user_post',methods=['POST','GET'])
def user_post():
    msg=""
    uname=""
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_user where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT * FROM mf_post where uname=%s order by id desc",(uname,))
    data2 = mycursor.fetchall()
    
    
    return render_template('web/user_post.html',msg=msg,data=data,data2=data2)

@app.route('/user_donate',methods=['POST','GET'])
def user_donate():
    msg=""
    uname=""
    st=""
    pid=request.args.get("pid")
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_user where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT count(*) FROM mf_donation where pid=%s order by id desc",(pid,))
    d1 = mycursor.fetchone()[0]
    if d1>0:
        st="1"
    
    mycursor.execute("SELECT * FROM mf_donation where pid=%s order by id desc",(pid,))
    data2 = mycursor.fetchall()
    
    
    return render_template('web/user_donate.html',msg=msg,data=data,data2=data2,st=st)

@app.route('/user_fund',methods=['POST','GET'])
def user_fund():
    msg=""
    uname=""
    st=""
    pid=request.args.get("pid")
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_user where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT count(*) FROM mf_donation where uname=%s order by id desc",(uname,))
    d1 = mycursor.fetchone()[0]
    if d1>0:
        st="1"
    
    mycursor.execute("SELECT * FROM mf_donation where uname=%s order by id desc",(uname,))
    data2 = mycursor.fetchall()
    
    
    return render_template('web/user_fund.html',msg=msg,data=data,data2=data2,st=st)

@app.route('/dr_home',methods=['POST','GET'])
def dr_home():
    msg=""
    uname=""
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_donator where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT * FROM mf_post p,mf_user u where p.uname=u.uname && p.req_status=1 && p.fund_amount=0 order by p.id desc")
    data2 = mycursor.fetchall()
    
    
    return render_template('web/dr_home.html',msg=msg,data=data,data2=data2)

@app.route('/dr_donate',methods=['POST','GET'])
def dr_donate():
    msg=""
    uname=""
    st=""
    pid=request.args.get("pid")
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_donator where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT * FROM mf_post where id=%s",(pid,))
    data2 = mycursor.fetchone()
    user=data2[1]
    famt=data2[17]

    now = datetime.datetime.now()
    rdate=now.strftime("%d-%m-%Y")

    if request.method=='POST':
        amount=request.form['amount']
        trans_id = request.form['trans_id']

        mycursor.execute("SELECT max(id)+1 FROM mf_donation")
        maxid = mycursor.fetchone()[0]
        if maxid is None:
            maxid=1
            
        amt=famt+int(amount)
        mycursor.execute("update mf_post set fund_amount=%s where id=%s",(amt,pid))
        mydb.commit()
        
         
        sql = "INSERT INTO mf_donation(id,donator,uname,pid,amount,trans_id,rdate) VALUES (%s, %s, %s, %s, %s,%s,%s)"
        val = (maxid,uname,user,pid,amount,trans_id,rdate)
        mycursor.execute(sql, val)
        mydb.commit()
        msg="success"

    mycursor.execute("SELECT count(*) FROM mf_donation where pid=%s",(pid,))
    d1 = mycursor.fetchone()[0]
    if d1>0:
        st="1"
    
    mycursor.execute("SELECT * FROM mf_donation where pid=%s",(pid,))
    data2 = mycursor.fetchall()
    
    return render_template('web/dr_donate.html',msg=msg,data=data,data2=data2,st=st)


@app.route('/dr_post',methods=['POST','GET'])
def dr_post():
    msg=""
    uname=""
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_donator where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT * FROM mf_post where req_status=1 order by id desc")
    data2 = mycursor.fetchall()
    
    
    return render_template('web/dr_post.html',msg=msg,data=data,data2=data2)

@app.route('/dr_bill', methods=['GET', 'POST'])
def dr_bill():
    msg=""
    st=""
    data2=[]
    view=""
    fname=""
    textpdf=""
    textdoc=""
    ccode=""
    pid=request.args.get("pid")
    nn=0
   
    ss='0'
    if pid=='0':
        ss='0'
    else:
        ss='1'
        
    mycursor = mydb.cursor()
 

    mycursor.execute("SELECT * FROM mf_files where post_id=%s",(pid,))
    data1 = mycursor.fetchall()

    mycursor.execute("SELECT count(*) FROM mf_files where post_id=%s",(pid,))
    cnt = mycursor.fetchone()[0]
    if cnt>0:
        st="1"
        
        mycursor.execute("SELECT * FROM mf_files where post_id=%s",(pid,))
        d1 = mycursor.fetchall()
        for d2 in d1:
            dt=[]
            dt.append(d2[0])
            dt.append(d2[1])
            dt.append(d2[2])
            dt.append(d2[3])
            dt.append(d2[4])
            dt.append(d2[5])
            dt.append(d2[6])

            #7#
            ex=d2[3].split(".")
            if ex[1]=="png":
                dt.append("png")
            elif ex[1]=="jpg":
                dt.append("jpg")
            elif ex[1]=="jpeg":
                dt.append("jpeg")
            elif ex[1]=="pdf":
                dt.append("pdf")
                fname=d2[3]
                file2="static/upload/"+fname
                textpdf = extract_text_from_pdf(file2)
               
            
            elif ex[1]=="docx":
                dt.append("docx")
                fname=d2[3]
                file2="static/upload/"+fname
    
                textdoc = extract_text_from_docx(file2)
                
            
            elif ex[1]=="txt":
                dt.append("txt")
                fname=d2[3]
                file1 = open("static/upload/"+fname, 'r')
                Lines = file1.readlines()
                 
                count = 0
                result=""
                # Strips the newline character
                for line in Lines:
                    result = "".join(line for line in Lines if not line.isspace())
                    count += 1
                    #print("Line{}: {}".format(count, line.strip()))
                ccode=result
            
            else:
                dt.append("")

            #8#
            if d2[6]>0:
                dt.append("yes")
            else:
                dt.append("no")
            #9#
            if d2[6]>0:
                i=1
                dtt=[]
                
                while i<=d2[6]:
                    mg="m"+str(d2[0])+"_"+str(i)+".png"
                    dtt.append(mg)
                    i+=1
                dt.append(dtt)
            ##
            data2.append(dt)
           

    return render_template('web/dr_bill.html',msg=msg,data2=data2,st=st)

@app.route('/dr_fund',methods=['POST','GET'])
def dr_fund():
    msg=""
    uname=""
    st=""
    if 'username' in session:
        uname = session['username']
        
    mycursor = mydb.cursor()
    mycursor.execute("SELECT * FROM mf_donator where uname=%s",(uname,))
    data = mycursor.fetchone()

    mycursor.execute("SELECT count(*) FROM mf_donation where donator=%s order by id desc",(uname,))
    d1 = mycursor.fetchone()[0]
    if d1>0:
        st="1"
    
    mycursor.execute("SELECT * FROM mf_donation where donator=%s order by id desc",(uname,))
    data2 = mycursor.fetchall()
    
    
    return render_template('web/dr_fund.html',msg=msg,data=data,data2=data2,st=st)


@app.route('/logout')
def logout():
    # remove the username from the session if it is there
    #session.pop('username', None)
    return redirect(url_for('index'))

import os
if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)